"""
app/services/docx_service.py
============================
Genera un .docx de cotización pre-llenado a partir de los datos del cliente.

Flujo:
  1. El endpoint POST /cotizacion llama a generar_cotizacion(payload, fila)
  2. Se copia la plantilla y se llenan los campos
  3. El archivo se guarda en DOCX_DIR / "COT-{fila}.docx"
  4. El admin descarga vía GET /api/v1/cotizacion/{fila}/docx

Dependencias:
  pip install python-docx
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.models.cotizacion import CotizacionRequest

# ── Rutas ─────────────────────────────────────────────────────────────────────
BASE_DIR  = Path(__file__).parent.parent.parent      # raíz del proyecto
TEMPLATE  = BASE_DIR / "plantilla_cotizacion.docx"   # plantilla original
DOCX_DIR  = BASE_DIR / "cotizaciones_generadas"      # carpeta de salida
DOCX_DIR.mkdir(exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _unique_cells(row):
    """Devuelve solo las celdas únicas de una fila (ignora merged duplicates)."""
    seen, cells = set(), []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            cells.append(cell)
    return cells


def _write(cell, text: str):
    """
    Escribe texto en una celda preservando el formato XML del primer run.
    Si la celda está vacía, crea un run nuevo.
    """
    text = str(text) if text is not None else ""
    para = cell.paragraphs[0]

    all_runs = [r for p in cell.paragraphs for r in p.runs]

    if all_runs:
        first_run = all_runs[0]
        for r in all_runs[1:]:
            r.text = ""
        first_run.text = text
    else:
        para.add_run(text)


def _fmt(value: float) -> str:
    """Formatea un número como moneda con separador de miles y 2 decimales."""
    return f"{value:,.2f}"


def numero_cotizacion(fila: int) -> str:
    year = date.today().year
    return f"COT-{fila:04d}-{year}"


# ─────────────────────────────────────────────────────────────────────────────
# FUNCIÓN PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────

def generar_cotizacion(data: CotizacionRequest, fila: int) -> Path:
    """
    Llena la plantilla .docx con los datos del cliente y la guarda en disco.

    Mapeo Tabla 1 — Datos del cliente:
      Fila 1, col 1 → Nombre                    (data.nombre)
      Fila 2, col 1 → Dirección                 (dirección del mapa)
      Fila 3, col 1 → Teléfono                  (data.telefono)
      Fila 3, col 3 → Correo                    (data.correo)
      Fila 4, col 1 → RTN                       (data.rtn)
      Fila 4, col 3 → Fecha de solicitud        (hoy)
      Fila 5, col 1 → Nombre del proyecto       (data.nombreProyecto)
      Fila 6, col 1 → Ubicación del proyecto    (data.direccionProyecto)
      Fila 7, col 1 → Contacto                  (correo, teléfono)

    Columnas de la Tabla 2 (unique cells por fila de datos):
      0 = No.
      1 = Descripción
      2 = Norma
      3 = No. Muestras
      4 = P. Unitario   ← precio unitario del servicio
      5 = Total         ← muestras × precio, calculado en Python

    Filas especiales:
      12 → SUBTOTAL  ← suma de todos los totales, calculada en Python
      13 → TOTAL     ← mismo valor que subtotal (el admin ajusta si aplica)

    Retorna el Path del archivo generado.
    """
    numero = numero_cotizacion(fila)
    doc    = Document(str(TEMPLATE))
    tables = doc.tables

    # ── TABLA 0 — Número de cotización ────────────────────────────────────────
    cot_cell = tables[0].rows[0].cells[1]
    if len(cot_cell.paragraphs) >= 2:
        num_para = cot_cell.paragraphs[1]
        if num_para.runs:
            num_para.runs[0].text = numero
        else:
            num_para.add_run(numero)

    # ── TABLA 1 — Datos del cliente ───────────────────────────────────────────
    t1 = tables[1]

    def fill_t1(row_idx: int, col_idx: int, value: str):
        unique = _unique_cells(t1.rows[row_idx])
        if col_idx < len(unique):
            _write(unique[col_idx], value)

    hoy = date.today().strftime("%d/%m/%Y")

    # Dirección desde el mapa (geocodificación inversa)
    direccion_mapa = ""
    if data.ubicacion:
        direccion_mapa = data.ubicacion.address or f"{data.ubicacion.lat}, {data.ubicacion.lng}"

    # Contacto: correo y teléfono separados por coma
    contacto = f"{data.correo}, {data.telefono}"

    fill_t1(1, 1, data.nombre)                          # Nombre del cliente
    fill_t1(2, 1, direccion_mapa)                       # Dirección (mapa)
    fill_t1(3, 1, data.telefono)                        # Teléfono
    fill_t1(3, 3, data.correo)                          # Correo
    fill_t1(4, 1, data.rtn or "")                       # RTN (del formulario)
    fill_t1(4, 3, hoy)                                  # Fecha de solicitud
    fill_t1(5, 1, data.nombreProyecto)                  # Nombre del proyecto
    fill_t1(6, 1, data.direccionProyecto or "")         # Ubicación del proyecto
    fill_t1(7, 1, contacto)                             # Contacto (correo, teléfono)

    # ── TABLA 2 — Detalle de ensayos ──────────────────────────────────────────
    t2         = tables[2]
    DATA_START = 2      # índice de la primera fila de datos
    MAX_ROWS   = 10

    servicios = data.servicios[:MAX_ROWS]

    # ── Llenar filas y acumular subtotal ──────────────────────────────────────
    subtotal = 0.0

    for i, svc in enumerate(servicios):
        muestras = svc.muestras or 0
        precio   = float(getattr(svc, 'precio', None) or 0)
        total    = muestras * precio
        subtotal += total

        row    = t2.rows[DATA_START + i]
        unique = _unique_cells(row)

        descripcion = svc.name + (f"\n({svc.sub})" if getattr(svc, 'sub', None) else "")

        if len(unique) > 1: _write(unique[1], descripcion)
        if len(unique) > 2: _write(unique[2], svc.norma or "")
        if len(unique) > 3: _write(unique[3], str(muestras) if muestras else "")
        if len(unique) > 4: _write(unique[4], _fmt(precio))
        if len(unique) > 5: _write(unique[5], _fmt(total))

    # Limpiar filas sobrantes (sin servicio)
    for i in range(len(servicios), MAX_ROWS):
        row    = t2.rows[DATA_START + i]
        unique = _unique_cells(row)
        for col in [1, 2, 3, 4, 5]:
            if col < len(unique):
                _write(unique[col], "")

    # ── SUBTOTAL — fila 12 ────────────────────────────────────────────────────
    row_sub    = t2.rows[12]
    unique_sub = _unique_cells(row_sub)
    if len(unique_sub) > 1:
        _write(unique_sub[-1], _fmt(subtotal))

    # ── TOTAL GENERAL — fila 13 ───────────────────────────────────────────────
    # Igual al subtotal; el admin puede ajustar manualmente si aplica descuento
    row_tot    = t2.rows[13]
    unique_tot = _unique_cells(row_tot)
    if len(unique_tot) > 1:
        _write(unique_tot[-1], _fmt(subtotal))

    # ── Guardar ───────────────────────────────────────────────────────────────
    output_path = DOCX_DIR / f"{numero}.docx"
    doc.save(str(output_path))
    return output_path